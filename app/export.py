"""
Экспорт результатов прогона графа.

to_docx(state, path)       — бизнес-отчёт (.docx): ранжирование гипотез,
                              обоснование оценок, источники (evidence).
to_pdf(state, path)        — тот же отчёт в PDF (reportlab).
to_tasks_csv(state, path)  — задачи по протоколам проверки (state["roadmap"])
                              в CSV, пригодном для импорта в Jira/YouTrack.
to_tasks_json(state, path) — то же самое в JSON (для импорта через API).

Оба отчёта строятся по одному и тому же контенту (ranked + roadmap), просто
в разных форматах — чтобы не разъезжалась логика между DOCX и PDF, порядок
секций и данные одинаковы.
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

import pandas as pd
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.state import State

logger = logging.getLogger("export")

# Шрифт с кириллицей для PDF: встроенные шрифты reportlab (Helvetica и т.п.)
# кириллицу не поддерживают. Шрифт вшит в репозиторий (app/assets/fonts),
# чтобы отчёт корректно собирался независимо от того, установлен ли DejaVu
# Sans в системе, где запускается прототип.
_FONT_DIR = Path(__file__).resolve().parent / "assets" / "fonts"
_FONT_REGULAR_PATH = _FONT_DIR / "DejaVuSans.ttf"
_FONT_BOLD_PATH = _FONT_DIR / "DejaVuSans-Bold.ttf"


def _register_pdf_fonts() -> tuple[str, str]:
    """Регистрирует DejaVu Sans в reportlab; при неудаче откатывается на Helvetica."""
    try:
        if "DejaVuSans" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("DejaVuSans", str(_FONT_REGULAR_PATH)))
            pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", str(_FONT_BOLD_PATH)))
        return "DejaVuSans", "DejaVuSans-Bold"
    except Exception as exc:
        logger.warning("Не удалось зарегистрировать DejaVu Sans (%s), использую Helvetica", exc)
        return "Helvetica", "Helvetica-Bold"


def _source_lookup(state: State) -> dict[str, str]:
    """source_id -> человекочитаемый source (файл + страница/лист) из retrieved."""
    return {c.source_id: c.source for c in state["retrieved"]}


def _fmt_dict(d: dict) -> str:
    return ", ".join(f"{k}: {v}" for k, v in d.items()) if d else "—"


def _fmt_list(items: list[str]) -> str:
    return "; ".join(items) if items else "—"


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def to_docx(state: State, path: str) -> None:
    """Собирает бизнес-отчёт (.docx): ранжированные гипотезы + обоснование + источники."""
    source_lookup = _source_lookup(state)
    doc = Document()

    doc.add_heading("Отчёт по гипотезам", level=0)
    doc.add_paragraph(f"Задача: {state['query']}")
    doc.add_paragraph(f"Ограничения: {_fmt_dict(state['constraints'])}")
    doc.add_paragraph(f"Одобрено гипотез: {len(state['approved'])} из {len(state['ranked'])}")

    for position, ranked in enumerate(state["ranked"], start=1):
        h = ranked.hypothesis
        mark = " [ОДОБРЕНО]" if h.id in state["approved"] else ""
        doc.add_heading(f"#{position}. {h.id} — score={ranked.score:.4f}{mark}", level=1)
        doc.add_paragraph(h.statement)

        doc.add_heading("Оценка", level=2)
        score_table = doc.add_table(rows=1, cols=6)
        score_table.style = "Table Grid"
        for cell, name in zip(score_table.rows[0].cells, ["novelty", "value", "feasibility", "risk", "cost_of_error", "score"]):
            cell.text = name
        values_row = score_table.add_row().cells
        for cell, value in zip(values_row, [h.novelty, h.value, h.feasibility, h.risk, h.cost_of_error, ranked.score]):
            cell.text = f"{value:.3f}"
        doc.add_paragraph(h.rationale)

        doc.add_heading("Как сформулировано", level=2)
        doc.add_paragraph(f"Метод: {h.derivation_method}")
        for i, step in enumerate(h.reasoning_steps, start=1):
            doc.add_paragraph(f"{i}. {step}", style="List Number")

        doc.add_heading("Источники (evidence)", level=2)
        evidence_table = doc.add_table(rows=1, cols=4)
        evidence_table.style = "Table Grid"
        for cell, name in zip(evidence_table.rows[0].cells, ["source_id", "источник", "факт", "как использован"]):
            cell.text = name
        for e in h.evidence:
            row = evidence_table.add_row().cells
            row[0].text = e.source_id
            row[1].text = source_lookup.get(e.source_id, "—")
            row[2].text = e.fact
            row[3].text = e.how_used

        doc.add_heading("Почему именно это", level=2)
        doc.add_paragraph(f"KPI: {h.kpi_link}")
        doc.add_paragraph(f"Новизна: {h.novelty_justification}")
        doc.add_paragraph("Отклонённые альтернативы:")
        for alt in h.rejected_alternatives:
            doc.add_paragraph(alt, style="List Bullet")

        doc.add_heading("При каких условиях", level=2)
        doc.add_paragraph(f"Условия: {_fmt_dict(h.conditions)}")
        doc.add_paragraph(f"Предположения: {_fmt_list(h.assumptions)}")
        doc.add_paragraph(f"Пределы применимости: {h.validity_limits}")

        doc.add_page_break()

    if state["roadmap"]:
        doc.add_heading("Протоколы проверки одобренных гипотез", level=1)
        for roadmap in state["roadmap"]:
            doc.add_heading(f"Протокол для {roadmap.hypothesis_id}", level=2)
            for i, step in enumerate(roadmap.steps, start=1):
                doc.add_paragraph(f"{i}. {step.name}", style="List Number")
                doc.add_paragraph(f"Ресурсы: {_fmt_list(step.resources)}")
                doc.add_paragraph(f"Критерии успеха: {_fmt_list(step.success_criteria)}")
                doc.add_paragraph(f"Критерии провала: {_fmt_list(step.failure_criteria)}")

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def to_pdf(state: State, path: str) -> None:
    """Тот же отчёт, что и to_docx, но в PDF (reportlab), с шрифтом-кириллицей."""
    font_normal, font_bold = _register_pdf_fonts()
    source_lookup = _source_lookup(state)

    base_styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleRu", parent=base_styles["Title"], fontName=font_bold, fontSize=20)
    h1 = ParagraphStyle("H1Ru", parent=base_styles["Heading1"], fontName=font_bold, fontSize=15)
    h2 = ParagraphStyle("H2Ru", parent=base_styles["Heading2"], fontName=font_bold, fontSize=12)
    normal = ParagraphStyle("NormalRu", parent=base_styles["Normal"], fontName=font_normal, fontSize=10, leading=14)

    def p(text) -> Paragraph:
        return Paragraph(xml_escape(str(text)), normal)

    story = [
        Paragraph("Отчёт по гипотезам", title_style),
        Spacer(1, 12),
        p(f"Задача: {state['query']}"),
        p(f"Ограничения: {_fmt_dict(state['constraints'])}"),
        p(f"Одобрено гипотез: {len(state['approved'])} из {len(state['ranked'])}"),
        Spacer(1, 12),
    ]

    table_style = TableStyle(
        [
            ("FONTNAME", (0, 0), (-1, -1), font_normal),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]
    )

    for position, ranked in enumerate(state["ranked"], start=1):
        h = ranked.hypothesis
        mark = " [ОДОБРЕНО]" if h.id in state["approved"] else ""
        story.append(Paragraph(xml_escape(f"#{position}. {h.id} — score={ranked.score:.4f}{mark}"), h1))
        story.append(p(h.statement))

        story.append(Paragraph("Оценка", h2))
        score_data = [
            ["novelty", "value", "feasibility", "risk", "cost_of_error", "score"],
            [f"{h.novelty:.3f}", f"{h.value:.3f}", f"{h.feasibility:.3f}", f"{h.risk:.3f}", f"{h.cost_of_error:.3f}", f"{ranked.score:.4f}"],
        ]
        story.append(Table(score_data, hAlign="LEFT", style=table_style))
        story.append(p(h.rationale))

        story.append(Paragraph("Как сформулировано", h2))
        story.append(p(f"Метод: {h.derivation_method}"))
        for i, step in enumerate(h.reasoning_steps, start=1):
            story.append(p(f"{i}. {step}"))

        story.append(Paragraph("Источники (evidence)", h2))
        evidence_data = [["source_id", "источник", "факт", "как использован"]]
        for e in h.evidence:
            evidence_data.append(
                [
                    xml_escape(e.source_id),
                    xml_escape(source_lookup.get(e.source_id, "—")),
                    Paragraph(xml_escape(e.fact), normal),
                    Paragraph(xml_escape(e.how_used), normal),
                ]
            )
        evidence_table = Table(evidence_data, hAlign="LEFT", colWidths=[2.2 * cm, 3.3 * cm, 5.2 * cm, 5.2 * cm])
        evidence_table.setStyle(table_style)
        story.append(evidence_table)

        story.append(Paragraph("Почему именно это", h2))
        story.append(p(f"KPI: {h.kpi_link}"))
        story.append(p(f"Новизна: {h.novelty_justification}"))
        for alt in h.rejected_alternatives:
            story.append(p(f"— {alt}"))

        story.append(Paragraph("При каких условиях", h2))
        story.append(p(f"Условия: {_fmt_dict(h.conditions)}"))
        story.append(p(f"Предположения: {_fmt_list(h.assumptions)}"))
        story.append(p(f"Пределы применимости: {h.validity_limits}"))

        story.append(PageBreak())

    if state["roadmap"]:
        story.append(Paragraph("Протоколы проверки одобренных гипотез", h1))
        for roadmap in state["roadmap"]:
            story.append(Paragraph(xml_escape(f"Протокол для {roadmap.hypothesis_id}"), h2))
            for i, step in enumerate(roadmap.steps, start=1):
                story.append(p(f"{i}. {step.name}"))
                story.append(p(f"Ресурсы: {_fmt_list(step.resources)}"))
                story.append(p(f"Критерии успеха: {_fmt_list(step.success_criteria)}"))
                story.append(p(f"Критерии провала: {_fmt_list(step.failure_criteria)}"))
                story.append(Spacer(1, 6))

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    SimpleDocTemplate(path, pagesize=A4).build(story)


# ---------------------------------------------------------------------------
# Задачи по протоколам (CSV / JSON) — для импорта в Jira/YouTrack
# ---------------------------------------------------------------------------

_TASK_COLUMNS = [
    "hypothesis_id",
    "step_index",
    "summary",
    "description",
    "resources",
    "success_criteria",
    "failure_criteria",
]


def _build_tasks(state: State) -> list[dict]:
    """Разворачивает state["roadmap"] в плоский список задач — одна задача на шаг протокола."""
    tasks = []
    for roadmap in state["roadmap"]:
        for i, step in enumerate(roadmap.steps, start=1):
            tasks.append(
                {
                    "hypothesis_id": roadmap.hypothesis_id,
                    "step_index": i,
                    "summary": f"[{roadmap.hypothesis_id}] Шаг {i}: {step.name}",
                    "description": (
                        f"Ресурсы: {_fmt_list(step.resources)}\n"
                        f"Критерии успеха: {_fmt_list(step.success_criteria)}\n"
                        f"Критерии провала: {_fmt_list(step.failure_criteria)}"
                    ),
                    "resources": _fmt_list(step.resources),
                    "success_criteria": _fmt_list(step.success_criteria),
                    "failure_criteria": _fmt_list(step.failure_criteria),
                }
            )
    return tasks


def to_tasks_csv(state: State, path: str) -> None:
    """
    Задачи по протоколам проверки в CSV — одна строка на шаг протокола, колонки
    summary/description совместимы с полями импорта Jira; остальные колонки
    (resources/success_criteria/failure_criteria) можно замэппить в кастомные
    поля YouTrack при импорте.
    """
    tasks = _build_tasks(state)
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(tasks, columns=_TASK_COLUMNS).to_csv(path, index=False)


def to_tasks_json(state: State, path: str) -> None:
    """То же самое в JSON — для импорта через API Jira/YouTrack или скриптом."""
    tasks = _build_tasks(state)
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(tasks, f, ensure_ascii=False, indent=2)
